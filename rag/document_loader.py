"""Document loader for the DataIntern RAG Engine.

Provides parsers for CSV, TSV, JSON, Excel, PDF, and DOCX files.
Each parser returns a list of `Document` dataclass instances with
structured text and rich metadata ready for downstream chunking and
embedding.
"""

from __future__ import annotations

import json
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import fitz  # PyMuPDF
import pandas as pd
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Data class
# ---------------------------------------------------------------------------


@dataclass
class Document:
    """A single unit of text extracted from a source file.

    Attributes:
        text: The extracted / formatted text content.
        metadata: A dictionary carrying provenance information such as
            filename, sheet, page, row_range, paragraph, chunk_id, and
            source_type.
    """

    text: str
    metadata: dict[str, Any] = field(default_factory=dict)


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

_TABULAR_ROW_GROUP_SIZE = 15  # rows per Document for CSV / TSV / Excel
_JSON_OBJECT_GROUP_SIZE = 10  # objects per Document for JSON arrays
_DOCX_PARA_GROUP_SIZE = 10   # paragraphs per Document for DOCX


# ---------------------------------------------------------------------------
# DocumentLoader
# ---------------------------------------------------------------------------


class DocumentLoader:
    """Discovers and parses files into a flat list of `Document` objects.

    Supported extensions: .csv, .tsv, .json, .xlsx, .xls, .pdf, .docx

    Usage::

        loader = DocumentLoader()
        docs = loader.load_all("./data", extensions=(".csv", ".json", ".xlsx"))
    """

    # ----- CSV / TSV helpers ------------------------------------------------

    @staticmethod
    def _format_tabular_rows(
        df: pd.DataFrame,
        filename: str,
        start_row: int,
        end_row: int,
    ) -> str:
        """Convert a DataFrame slice into a structured key-value text block.

        Args:
            df: The slice of the dataframe to format.
            filename: Source filename for the header line.
            start_row: 1-based start row index.
            end_row: 1-based end row index (inclusive).

        Returns:
            A human-readable text block with a header line and one pipe-
            delimited row per record.
        """
        header = f"File: {filename} | Rows: {start_row}-{end_row}"
        columns = list(df.columns)
        lines: list[str] = [header]
        for _, row in df.iterrows():
            parts = [f"{col}: {row[col]}" for col in columns]
            lines.append(" | ".join(parts))
        return "\n".join(lines)

    # ----- Parsers ----------------------------------------------------------

    def parse_csv(self, filepath: Path) -> list[Document]:
        """Parse a CSV file into grouped Documents.

        Every *_TABULAR_ROW_GROUP_SIZE* rows are combined into a single
        Document whose text includes the column headers and formatted
        key-value rows.

        Args:
            filepath: Path to the CSV file.

        Returns:
            A list of Document instances.
        """
        logger.info("Parsing CSV: %s", filepath.name)
        df = pd.read_csv(filepath)
        return self._dataframe_to_documents(df, filepath.name, source_type="csv")

    def parse_tsv(self, filepath: Path) -> list[Document]:
        """Parse a TSV file into grouped Documents.

        Identical to :meth:`parse_csv` but uses a tab separator.

        Args:
            filepath: Path to the TSV file.

        Returns:
            A list of Document instances.
        """
        logger.info("Parsing TSV: %s", filepath.name)
        df = pd.read_csv(filepath, sep="\t")
        return self._dataframe_to_documents(df, filepath.name, source_type="tsv")

    def parse_json(self, filepath: Path) -> list[Document]:
        """Parse a JSON file into grouped Documents.

        Handles two layouts:
        1. **Top-level array** – ``[{...}, {...}, ...]``
        2. **Nested wrapper** – ``{"key": [{...}, ...], ...}``  where the
           first key whose value is a list is used.

        Objects are grouped in batches of *_JSON_OBJECT_GROUP_SIZE* and
        converted to readable key-value text.

        Args:
            filepath: Path to the JSON file.

        Returns:
            A list of Document instances.
        """
        logger.info("Parsing JSON: %s", filepath.name)
        with open(filepath, "r", encoding="utf-8") as fh:
            data = json.load(fh)

        # Normalise to a flat list of dicts
        records: list[dict[str, Any]] = []
        if isinstance(data, list):
            records = data
        elif isinstance(data, dict):
            # Find the first key that holds a list of dicts
            for key, value in data.items():
                if isinstance(value, list) and value and isinstance(value[0], dict):
                    records = value
                    logger.debug(
                        "JSON wrapper detected – using key '%s' (%d records)",
                        key,
                        len(value),
                    )
                    break
            if not records:
                # Fallback: treat the whole dict as one record
                records = [data]

        docs: list[Document] = []
        group_size = _JSON_OBJECT_GROUP_SIZE
        for i in range(0, len(records), group_size):
            group = records[i : i + group_size]
            start_idx = i + 1
            end_idx = i + len(group)
            header = f"File: {filepath.name} | Records: {start_idx}-{end_idx}"
            lines: list[str] = [header]
            for obj in group:
                parts = [f"{k}: {v}" for k, v in obj.items()]
                lines.append(" | ".join(parts))
            docs.append(
                Document(
                    text="\n".join(lines),
                    metadata={
                        "filename": filepath.name,
                        "record_range": f"{start_idx}-{end_idx}",
                        "source_type": "json",
                    },
                )
            )

        logger.info("  → %d document(s) from %s", len(docs), filepath.name)
        return docs

    def parse_excel(self, filepath: Path) -> list[Document]:
        """Parse an Excel workbook (.xlsx / .xls) into grouped Documents.

        All sheets are parsed.  Each sheet's rows are grouped identically
        to the CSV parser, and the sheet name is recorded in metadata.

        Args:
            filepath: Path to the Excel file.

        Returns:
            A list of Document instances covering every sheet.
        """
        logger.info("Parsing Excel: %s", filepath.name)
        xls = pd.ExcelFile(filepath)
        all_docs: list[Document] = []
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            sheet_docs = self._dataframe_to_documents(
                df,
                filepath.name,
                source_type="excel",
                sheet=sheet_name,
            )
            all_docs.extend(sheet_docs)
            logger.debug(
                "  Sheet '%s': %d document(s)", sheet_name, len(sheet_docs)
            )
        logger.info("  → %d document(s) from %s", len(all_docs), filepath.name)
        return all_docs

    def parse_pdf(self, filepath: Path) -> list[Document]:
        """Parse a PDF file into per-page Documents.

        Uses PyMuPDF (``fitz``) to extract text.  Each page produces one
        Document.

        Args:
            filepath: Path to the PDF file.

        Returns:
            A list of Document instances (one per page).
        """
        logger.info("Parsing PDF: %s", filepath.name)
        docs: list[Document] = []
        with fitz.open(filepath) as pdf:
            for page_num in range(len(pdf)):
                page = pdf[page_num]
                text = page.get_text()
                if text.strip():
                    docs.append(
                        Document(
                            text=f"File: {filepath.name} | Page: {page_num + 1}\n{text.strip()}",
                            metadata={
                                "filename": filepath.name,
                                "page": page_num + 1,
                                "source_type": "pdf",
                            },
                        )
                    )
        logger.info("  → %d document(s) from %s", len(docs), filepath.name)
        return docs

    def parse_docx(self, filepath: Path) -> list[Document]:
        """Parse a DOCX file into grouped paragraph Documents.

        Non-empty paragraphs are batched in groups of
        *_DOCX_PARA_GROUP_SIZE*.  Tables are extracted separately—each
        table becomes its own Document with pipe-delimited formatting.

        Args:
            filepath: Path to the DOCX file.

        Returns:
            A list of Document instances.
        """
        logger.info("Parsing DOCX: %s", filepath.name)
        doc = DocxDocument(str(filepath))
        docs: list[Document] = []

        # --- Paragraphs -----------------------------------------------------
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        group_size = _DOCX_PARA_GROUP_SIZE
        for i in range(0, len(paragraphs), group_size):
            group = paragraphs[i : i + group_size]
            start_idx = i + 1
            end_idx = i + len(group)
            header = f"File: {filepath.name} | Paragraphs: {start_idx}-{end_idx}"
            text = header + "\n" + "\n".join(group)
            docs.append(
                Document(
                    text=text,
                    metadata={
                        "filename": filepath.name,
                        "paragraph": f"{start_idx}-{end_idx}",
                        "source_type": "docx",
                    },
                )
            )

        # --- Tables ----------------------------------------------------------
        for table_idx, table in enumerate(doc.tables, start=1):
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                header = f"File: {filepath.name} | Table: {table_idx}"
                text = header + "\n" + "\n".join(rows)
                docs.append(
                    Document(
                        text=text,
                        metadata={
                            "filename": filepath.name,
                            "table": table_idx,
                            "source_type": "docx",
                        },
                    )
                )

        logger.info("  → %d document(s) from %s", len(docs), filepath.name)
        return docs

    # ----- Orchestrator -----------------------------------------------------

    def load_file(self, filepath: Path) -> list[Document]:
        """Parse a single file based on its extension."""
        if not filepath.is_file():
            raise FileNotFoundError(f"File not found: {filepath}")
            
        ext = filepath.suffix.lower()
        parser_map: dict[str, Any] = {
            ".csv": self.parse_csv,
            ".tsv": self.parse_tsv,
            ".json": self.parse_json,
            ".xlsx": self.parse_excel,
            ".xls": self.parse_excel,
            ".pdf": self.parse_pdf,
            ".docx": self.parse_docx,
        }
        
        parser = parser_map.get(ext)
        if parser is None:
            logger.warning("No parser for extension '%s' – skipping %s", ext, filepath.name)
            return []
            
        try:
            docs = parser(filepath)
            logger.info("  %s → %d document(s)", filepath.name, len(docs))
            return docs
        except Exception:
            logger.exception("Failed to parse %s", filepath.name)
            return []

    def load_all(
        self,
        data_dir: str,
        extensions: tuple[str, ...] = (
            ".csv",
            ".tsv",
            ".json",
            ".xlsx",
            ".xls",
            ".pdf",
            ".docx",
        ),
    ) -> list[Document]:
        """Discover and parse all supported files under *data_dir*.

        Each file is routed to the appropriate parser based on its
        extension.  Unsupported extensions are silently skipped.

        Args:
            data_dir: Root directory to scan for data files.
            extensions: Tuple of file extensions to include
                (with leading dot, lower-case).

        Returns:
            A flat list of all parsed Document instances.
        """
        data_path = Path(data_dir)
        if not data_path.is_dir():
            raise FileNotFoundError(f"Data directory not found: {data_dir}")

        parser_map: dict[str, Any] = {
            ".csv": self.parse_csv,
            ".tsv": self.parse_tsv,
            ".json": self.parse_json,
            ".xlsx": self.parse_excel,
            ".xls": self.parse_excel,
            ".pdf": self.parse_pdf,
            ".docx": self.parse_docx,
        }

        all_docs: list[Document] = []
        files = sorted(
            f for f in data_path.iterdir() if f.is_file() and f.suffix.lower() in extensions
        )

        logger.info(
            "Discovered %d file(s) in '%s' matching %s",
            len(files),
            data_dir,
            extensions,
        )

        for filepath in files:
            ext = filepath.suffix.lower()
            parser = parser_map.get(ext)
            if parser is None:
                logger.warning("No parser for extension '%s' – skipping %s", ext, filepath.name)
                continue
            try:
                docs = parser(filepath)
                all_docs.extend(docs)
                logger.info("  %s → %d document(s)", filepath.name, len(docs))
            except Exception:
                logger.exception("Failed to parse %s", filepath.name)

        logger.info("Total documents loaded: %d", len(all_docs))
        return all_docs

    # ----- Internal helpers -------------------------------------------------

    def _dataframe_to_documents(
        self,
        df: pd.DataFrame,
        filename: str,
        source_type: str,
        sheet: str | None = None,
    ) -> list[Document]:
        """Convert a pandas DataFrame into grouped Document instances.

        Args:
            df: The dataframe to convert.
            filename: Source filename for metadata.
            source_type: One of 'csv', 'tsv', 'excel'.
            sheet: Optional sheet name (Excel only).

        Returns:
            A list of Document instances.
        """
        docs: list[Document] = []
        total_rows = len(df)
        group_size = _TABULAR_ROW_GROUP_SIZE

        for start in range(0, total_rows, group_size):
            end = min(start + group_size, total_rows)
            chunk_df = df.iloc[start:end]
            start_row = start + 1  # 1-based
            end_row = end
            text = self._format_tabular_rows(chunk_df, filename, start_row, end_row)

            meta: dict[str, Any] = {
                "filename": filename,
                "row_range": f"{start_row}-{end_row}",
                "source_type": source_type,
            }
            if sheet is not None:
                meta["sheet"] = sheet

            docs.append(Document(text=text, metadata=meta))

        logger.info("  → %d document(s) from %s", len(docs), filename)
        return docs
